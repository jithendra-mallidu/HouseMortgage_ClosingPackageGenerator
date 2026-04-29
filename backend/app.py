"""
Closing Package Generator - Flask API
--------------------------------------
REST API backend for generating merged closing package PDFs by filling in
the actual underscore placeholders inside the Fannie Mae templates:
  1. State-specific Security Instrument (Mortgage / Deed of Trust / Security Deed)
  2. Promissory Note (state-specific if available, otherwise Multistate Form 3200)
  3. Notice of Right to Cancel

Deployed on Google Cloud Run.
"""
from __future__ import annotations

import copy
import io
import os
import random
import re
import shutil
import subprocess
import tempfile
from datetime import date, timedelta
from pathlib import Path

from flask import Flask, request, send_file, abort, jsonify
from flask_cors import CORS
from PyPDF2 import PdfReader, PdfWriter
from reportlab.lib.pagesizes import LETTER
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
from reportlab.pdfgen import canvas
from reportlab.lib import colors

from docx import Document as DocxDocument

# -----------------------------------------------------------------------------
# Paths — relative to this file for Docker deployment
# -----------------------------------------------------------------------------
BASE = Path(__file__).parent
TEMPLATE_DIR = BASE / "templates"
SI_DIR = TEMPLATE_DIR / "security_instruments"
NOTE_DIR = TEMPLATE_DIR / "notes"
RIDER_DIR = TEMPLATE_DIR / "riders"
SP_DIR = TEMPLATE_DIR / "special_purpose"
NOTICE_PDF = TEMPLATE_DIR / "Notice_Of_Right_To_Cancel.pdf"
CA_ACK_PDF = TEMPLATE_DIR / "Notary_Acknowledgment_California.pdf"
CACHE = BASE / "cache"
CACHE.mkdir(exist_ok=True)

# -----------------------------------------------------------------------------
# State registry: state name -> (security instrument docx, note pdf filename)
# If a state has no state-specific Note PDF, we fall back to the Multistate Note.
# -----------------------------------------------------------------------------
MULTISTATE_NOTE = "Form 3200 - Multistate Fixed Rate Note 07_2021(2081216.1)"

# Map of state -> Security Instrument DOCX filename
SECURITY_INSTRUMENTS = {
    "Alabama": "Form 3001 - Alabama Mortgage 07_2021.docx",
    "Alaska": "Form 3002 - Alaska Deed of Trust 07_2021.docx",
    "Arizona": "Form 3003 - Arizona Deed of Trust 07_2021.docx",
    "Arkansas": "Form 3004 - Arkansas Mortgage 07_2021.docx",
    "California": "Form 3005 - California Deed of Trust 07_2021.docx",
    "Colorado": "Form 3006 - Colorado Deed of Trust 07_2021.docx",
    "Connecticut": "Form 3007 - Connecticut Mortgage 07_2021 (rev. 07_24).docx",
    "Delaware": "Form 3008 - Delaware Mortgage 07_2021 (rev. 08_22).docx",
    "District of Columbia": "Form 3009 - DC Deed of Trust 07_2021 (rev. 08_23).docx",
    "Florida": "Form 3010 - Florida Mortgage 07_2021.docx",
    "Georgia": "Form 3011 - Georgia Security Deed 07_2021 (rev. 02_22).docx",
    "Hawaii": "Form 3012 - Hawaii Mortgage 07_2021.docx",
    "Idaho": "Form 3013 - Idaho Deed of Trust 07_2021.docx",
    "Illinois": "Form 3014 - Illinois Mortgage 07_2021.docx",
    "Indiana": "Form 3015 - Indiana Mortgage 07_2021.docx",
    "Iowa": "Form 3016 - Iowa Mortgage 07_2021.docx",
    "Kansas": "Form 3017 - Kansas Mortgage 07_2021.docx",
    "Kentucky": "Form 3018 - Kentucky Mortgage 07_2021 (rev. 07_24).docx",
    "Louisiana": "Form 3019 - Louisiana Mortgage 07_2021.docx",
    "Maine": "Form 3020 - Maine Mortgage 07_2021.docx",
    "Maryland": "Form 3021 - Maryland Deed of Trust 07_2021 (rev. 06_25).docx",
    "Massachusetts": "Form 3022 - Massachusetts Mortgage 07_2021.docx",
    "Michigan": "Form 3023 - Michigan Mortgage 07_2021.docx",
    "Minnesota": "Form 3024 - Minnesota Mortgage 07_2021.docx",
    "Mississippi": "Form 3025 - Mississippi Deed of Trust 07_2021.docx",
    "Missouri": "Form 3026 - Missouri Deed of Trust 07_2021.docx",
    "Montana": "Form 3027 - Montana Deed of Trust 07_2021.docx",
    "Nebraska": "Form 3028 - Nebraska Deed of Trust 07_2021.docx",
    "Nevada": "Form 3029 - Nevada Deed of Trust 07_2021.docx",
    "New Hampshire": "Form 3030 - New Hampshire Mortgage 07_2021.docx",
    "New Jersey": "Form 3031 - New Jersey Mortgage 07_2021.docx",
    "New Mexico": "Form 3032 - New Mexico Mortgage 07_2021.docx",
    "New York": "Form 3033 - New York Mortgage 07_2021.docx",
    "North Carolina": "Form 3034 - North Carolina Deed of Trust 07_2021.docx",
    "North Dakota": "Form 3035 - North Dakota Mortgage 07_2021.docx",
    "Ohio": "Form 3036 - Ohio Mortgage 07_2021.docx",
    "Oklahoma": "Form 3037 - Oklahoma Mortgage 07_2021.docx",
    "Oregon": "Form 3038 - Oregon Deed of Trust 07_2021.docx",
    "Pennsylvania": "Form 3039 - Pennsylvania Mortgage 07_2021.docx",
    "Rhode Island": "Form 3040 - Rhode Island Mortgage 07_2021.docx",
    "South Carolina": "Form 3041 - South Carolina Mortgage 07_2021.docx",
    "South Dakota": "Form 3042 - South Dakota Mortgage 07_2021.docx",
    "Tennessee": "Form 3043 - Tennessee Deed of Trust 07_2021 (rev. 09_25).docx",
    "Texas": "Form 3044 - Texas Deed of Trust 07_2021 (rev. 01_24).docx",
    "Utah": "Form 3045 - Utah Deed of Trust 07_2021.docx",
    "Vermont": "Form 3046 - Vermont Mortgage 07_2021.docx",
    "Virginia": "Form 3047 - Virginia Deed of Trust 07_2021 (rev. 02_23).docx",
    "Washington": "Form 3048 - Washington Deed of Trust 07_2021 (rev. 09_22).docx",
    "West Virginia": "Form 3049 - West Virginia Deed of Trust 07_2021.docx",
    "Wisconsin": "Form 3050 - Wisconsin Mortgage 07_2021.docx",
    "Wyoming": "Form 3051 - Wyoming Mortgage 07_2021.docx",
    "Guam": "Form 3052 - Guam Mortgage 07_2021.docx",
    "Puerto Rico": "Form 3053 - Puerto Rico Mortgage 07_2021 (rev. 03_23).docx",
    "US Virgin Islands": "Form 3054 - US Virgin Islands Mortgage 07_2021.docx",
}

# States that have a state-specific Note (otherwise we use the Multistate Note)
STATE_SPECIFIC_NOTES = {
    "Alaska": "Form 3202 - Alaska Fixed Rate Note 07_2021(2081222.1)",
    "Florida": "Form 3210 - Florida Fixed Rate Note 07_2021(2081225.1)",
    "Maine": "Form 3220 - Maine Fixed Rate Note 07_2021(2081232.1)",
    "New Hampshire": "Form 3230 - New Hampshire Fixed Rate Note 07_2021(2081234.1)",
    "New York": "Form 3233 - New York Fixed Rate Note 07_2021(2081236.1)",
    "Pennsylvania": "Form 3239 - Pennsylvania Fixed Rate Note 07_2021(2081237.1)",
    "Vermont": "Form 3246 - Vermont Fixed Rate Note 07_2021(2081241.1)",
    "Virginia": "Form 3247 - Virginia Fixed Rate Note 07_2021(2081242.1)",
    "West Virginia": "Form 3249 - West Virginia Fixed Rate Note 07_2021(2081243.1)",
    "Wisconsin": "Form 3250 - Wisconsin Fixed Rate Note 07_2021(2081251.1)",
    "Puerto Rico": "Form 3253 - Puerto Rico Fixed Rate Note 07_2021(2081256.1)",
}

STATE_ABBR = {
    "Alabama": "AL", "Alaska": "AK", "Arizona": "AZ", "Arkansas": "AR",
    "California": "CA", "Colorado": "CO", "Connecticut": "CT", "Delaware": "DE",
    "District of Columbia": "DC", "Florida": "FL", "Georgia": "GA", "Hawaii": "HI",
    "Idaho": "ID", "Illinois": "IL", "Indiana": "IN", "Iowa": "IA", "Kansas": "KS",
    "Kentucky": "KY", "Louisiana": "LA", "Maine": "ME", "Maryland": "MD",
    "Massachusetts": "MA", "Michigan": "MI", "Minnesota": "MN", "Mississippi": "MS",
    "Missouri": "MO", "Montana": "MT", "Nebraska": "NE", "Nevada": "NV",
    "New Hampshire": "NH", "New Jersey": "NJ", "New Mexico": "NM", "New York": "NY",
    "North Carolina": "NC", "North Dakota": "ND", "Ohio": "OH", "Oklahoma": "OK",
    "Oregon": "OR", "Pennsylvania": "PA", "Rhode Island": "RI", "South Carolina": "SC",
    "South Dakota": "SD", "Tennessee": "TN", "Texas": "TX", "Utah": "UT",
    "Vermont": "VT", "Virginia": "VA", "Washington": "WA", "West Virginia": "WV",
    "Wisconsin": "WI", "Wyoming": "WY", "Guam": "GU", "Puerto Rico": "PR",
    "US Virgin Islands": "VI",
}

# States that use a Deed of Trust (and therefore have a Trustee)
DEED_OF_TRUST_STATES = {
    "Alaska", "Arizona", "California", "Colorado", "District of Columbia",
    "Idaho", "Maryland", "Mississippi", "Missouri", "Montana", "Nebraska",
    "Nevada", "North Carolina", "Oregon", "Tennessee", "Texas", "Utah",
    "Virginia", "Washington", "West Virginia",
}

# -----------------------------------------------------------------------------
# Additional Documents Registry
# Each entry: id -> {name, category, filename, dir, states}
#   states: list of states this doc applies to, or None for all states
# -----------------------------------------------------------------------------
ADDITIONAL_DOCUMENTS = {
    # ── Riders ──
    "condo_rider": {
        "name": "Condominium Rider",
        "category": "Riders",
        "filename": "Form 3140 - Multistate Condominium Rider 07_2021.pdf",
        "dir": "riders",
        "states": None,  # all states
    },
    "condo_rider_tx": {
        "name": "Texas Home Equity Condominium Rider",
        "category": "Riders",
        "filename": "Form 3140.44 - Texas Home Equity Condo Rider 07_2021.pdf",
        "dir": "riders",
        "states": ["Texas"],
    },
    "condo_rider_pr": {
        "name": "Puerto Rico Condominium Rider",
        "category": "Riders",
        "filename": "Form 3140.53 - Puerto Rico Condominium Rider  07_2021.pdf",
        "dir": "riders",
        "states": ["Puerto Rico"],
    },
    "sofr_arm_rider": {
        "name": "SOFR ARM Rider",
        "category": "Riders",
        "filename": "Form 3141 - Multistate SOFR ARM Rider 07_2021.pdf",
        "dir": "riders",
        "states": None,
    },
    "sofr_arm_rider_wv": {
        "name": "West Virginia SOFR ARM Rider",
        "category": "Riders",
        "filename": "Form 3141.49 - West Virginia SOFR ARM Rider 05_2024.pdf",
        "dir": "riders",
        "states": ["West Virginia"],
    },
    "sofr_arm_rider_pr": {
        "name": "Puerto Rico SOFR ARM Rider",
        "category": "Riders",
        "filename": "Form 3141.53 - Puerto Rico SOFR ARM Rider 07_2021.pdf",
        "dir": "riders",
        "states": ["Puerto Rico"],
    },
    "fixed_adj_sofr_rider": {
        "name": "Fixed/Adjustable Rate Rider (30-Day Avg SOFR)",
        "category": "Riders",
        "filename": "Form 3142 - Multistate Fixed_Adjustable Rate Rider - 30 Day Average SOFR 07.2021.pdf",
        "dir": "riders",
        "states": None,
    },
    "fixed_adj_sofr_rider_tx": {
        "name": "Texas Home Equity Fixed/Adjustable SOFR Rider",
        "category": "Riders",
        "filename": "Form 3142.44 - Texas Home Equity Fixed Adjustable SOFR Rider 07_2021.pdf",
        "dir": "riders",
        "states": ["Texas"],
    },
    "fixed_adj_sofr_rider_wv": {
        "name": "West Virginia Fixed/Adjustable Rate Rider (SOFR)",
        "category": "Riders",
        "filename": "Form 3142.49 - West Virginia Fixed_Adjustable Rate Rider - 30 Day Average SOFR 05_2024.pdf",
        "dir": "riders",
        "states": ["West Virginia"],
    },
    "fixed_adj_sofr_rider_pr": {
        "name": "Puerto Rico Fixed/Adjustable ARM Rider (SOFR)",
        "category": "Riders",
        "filename": "Form 3142.53 - Puerto Rico SOFR F_A ARM Rider 07_2021.pdf",
        "dir": "riders",
        "states": ["Puerto Rico"],
    },
    "pud_rider": {
        "name": "PUD Rider",
        "category": "Riders",
        "filename": "Form 3150 - Multistate PUD Rider 07_2021.pdf",
        "dir": "riders",
        "states": None,
    },
    "pud_rider_tx": {
        "name": "Texas Home Equity PUD Rider",
        "category": "Riders",
        "filename": "Form 3150.44 - Texas Home Equity PUD Rider 07_2021.pdf",
        "dir": "riders",
        "states": ["Texas"],
    },
    "pud_rider_pr": {
        "name": "Puerto Rico PUD Rider",
        "category": "Riders",
        "filename": "Form 3150.53 - Puerto Rico PUD Rider 07_2021.pdf",
        "dir": "riders",
        "states": ["Puerto Rico"],
    },
    "mers_rider": {
        "name": "MERS Rider",
        "category": "Riders",
        "filename": "Form 3158 - MERs Rider 07_2021.pdf",
        "dir": "riders",
        "states": None,
    },
    "family_rider": {
        "name": "1-4 Family Rider",
        "category": "Riders",
        "filename": "Form 3170 - 1-4 Family Rider 07_2021.pdf",
        "dir": "riders",
        "states": None,
    },
    "family_rider_pr": {
        "name": "Puerto Rico 1-4 Family Rider",
        "category": "Riders",
        "filename": "Form 3170.53 - Puerto Rico 1-4 Family Rider 07_2021.pdf",
        "dir": "riders",
        "states": ["Puerto Rico"],
    },
    "second_home_rider": {
        "name": "Second Home Rider",
        "category": "Riders",
        "filename": "Form 3890 - Multistate Second Home Rider 07_2021.pdf",
        "dir": "riders",
        "states": None,
    },
    "second_home_rider_pr": {
        "name": "Puerto Rico Second Home Rider",
        "category": "Riders",
        "filename": "Form 3890.53 - Puerto Rico Second Home Rider 07_2021.pdf",
        "dir": "riders",
        "states": ["Puerto Rico"],
    },
    "reno_loan_rider": {
        "name": "Renovation Loan Rider to Security Instrument",
        "category": "Riders",
        "filename": "Form 3732 - Multistate Reno Loan Rider to SI FNMA Model Doc 07.2021.pdf",
        "dir": "riders",
        "states": None,
    },
    "investor_reno_rider": {
        "name": "Investor Renovation Loan Rider to Security Instrument",
        "category": "Riders",
        "filename": "Form 3733 - Multistate Investor Reno Loan Rider to SI FNMA Model Doc 07.2021.pdf",
        "dir": "riders",
        "states": None,
    },
    "construction_addendum_note": {
        "name": "Construction Loan Addendum to Note",
        "category": "Riders",
        "filename": "Form 3736 - Multistate Construction Loan Addendum to Note FNMA Model Doc 07.2021.pdf",
        "dir": "riders",
        "states": None,
    },
    "construction_rider": {
        "name": "Construction Loan Rider",
        "category": "Riders",
        "filename": "Form 3737 - Multistate Construction Loan Rider FNMA Model Doc 07_2021.pdf",
        "dir": "riders",
        "states": None,
    },
    "investor_construction_rider": {
        "name": "Investor Construction Loan Rider to Security Instrument",
        "category": "Riders",
        "filename": "Form 3738 - Multistate Investor Construction Loan Rider to SI FNMA Model Doc 07_2021.pdf",
        "dir": "riders",
        "states": None,
    },
    "pr_acceptance_mortgage": {
        "name": "Puerto Rico Acceptance of Mortgage",
        "category": "Riders",
        "filename": "Form 3053.1 - Puerto Rico Acceptance of Mortgage 07_2021.pdf",
        "dir": "riders",
        "states": ["Puerto Rico"],
    },
    # ── Special Purpose Documents ──
    "ny_cema": {
        "name": "New York CEMA",
        "category": "Special Purpose Documents",
        "filename": "Form 3172 - New York CEMA 07_2021.pdf",
        "dir": "special_purpose",
        "states": ["New York"],
    },
    "tx_home_equity_affidavit": {
        "name": "Texas Home Equity Affidavit Agreement",
        "category": "Special Purpose Documents",
        "filename": "Form 3185 - Texas_Home Equity Affidavit Agreement 07_2021.pdf",
        "dir": "special_purpose",
        "states": ["Texas"],
    },
    "renovation_contract": {
        "name": "Renovation Contract",
        "category": "Special Purpose Documents",
        "filename": "Form 3730 - Multistate Renovation Contract - Fannie Mae Model Document 07_2021 (rev. 11_23).pdf",
        "dir": "special_purpose",
        "states": None,
    },
    "renovation_loan_agreement": {
        "name": "Renovation Loan Agreement",
        "category": "Special Purpose Documents",
        "filename": "Form 3731 - Multistate Renovation Loan Agreement - Fannie Mae Model Document 07_2021 (rev. 02_22).pdf",
        "dir": "special_purpose",
        "states": None,
    },
    "construction_contract": {
        "name": "Construction Contract",
        "category": "Special Purpose Documents",
        "filename": "Form 3734 - Multistate Construction Contract - FNMA Model Document 07 2021.pdf",
        "dir": "special_purpose",
        "states": None,
    },
    "construction_loan_agreement": {
        "name": "Construction Loan Agreement",
        "category": "Special Purpose Documents",
        "filename": "Form 3735 - Multistate Construction Loan Agreement - FNMA Model Doc 07_2021.pdf",
        "dir": "special_purpose",
        "states": None,
    },
    "lien_waiver": {
        "name": "Lien Waiver",
        "category": "Special Purpose Documents",
        "filename": "Form 3739 - Multistate Lien Waiver 07_2021.pdf",
        "dir": "special_purpose",
        "states": None,
    },
    "mortgage_assignment_3741": {
        "name": "Mortgage Assignment (Form 3741)",
        "category": "Special Purpose Documents",
        "filename": "Form 3741 - Multistate Mortgage Assignment 07_2021.pdf",
        "dir": "special_purpose",
        "states": None,
    },
    "mortgage_assignment_3742": {
        "name": "Mortgage Assignment (Form 3742)",
        "category": "Special Purpose Documents",
        "filename": "Form 3742 - Multistate Mortgage Assignment 07_2021.pdf",
        "dir": "special_purpose",
        "states": None,
    },
    "dot_assignment_3743": {
        "name": "Deed of Trust Assignment (Form 3743)",
        "category": "Special Purpose Documents",
        "filename": "Form 3743 - Multistate Deed of Trust Assignment 07_2021.pdf",
        "dir": "special_purpose",
        "states": None,
    },
    "dot_assignment_3744": {
        "name": "Deed of Trust Assignment (Form 3744)",
        "category": "Special Purpose Documents",
        "filename": "Form 3744 - Multistate Deed of Trust Assignment 07_2021.pdf",
        "dir": "special_purpose",
        "states": None,
    },
    "ga_security_deed_assignment": {
        "name": "Georgia Security Deed Assignment",
        "category": "Special Purpose Documents",
        "filename": "Form 3745 - Georgia Security Deed Assignment 07_2021.pdf",
        "dir": "special_purpose",
        "states": ["Georgia"],
    },
    "la_notarial_endorsement": {
        "name": "Louisiana Notarial Endorsement Note Mortgage Assignment",
        "category": "Special Purpose Documents",
        "filename": "Form 3746 - Louisiana Notarial Endorsement Note Mortgage Assignment 07.2021.pdf",
        "dir": "special_purpose",
        "states": ["Louisiana"],
    },
    "subordination_refi": {
        "name": "Subordination Agreement (Refi Mortgage)",
        "category": "Special Purpose Documents",
        "filename": "Form 3747 - Subordination Agreement (Refi Mortgage) 07_2021.pdf",
        "dir": "special_purpose",
        "states": None,
    },
    "subordination_modified": {
        "name": "Subordination Agreement (Modified Mortgage)",
        "category": "Special Purpose Documents",
        "filename": "Form 3748 - Subordination Agreement (Modified Mortgage) 07_2021.pdf",
        "dir": "special_purpose",
        "states": None,
    },
    "mers_mortgage_assignment_me": {
        "name": "MERS Mortgage Assignment (Maine)",
        "category": "Special Purpose Documents",
        "filename": "Form 3749 - MERS Mortgage Assignment (Maine)  07_2021.pdf",
        "dir": "special_purpose",
        "states": ["Maine"],
    },
    "ms_ground_lease_renewals": {
        "name": "Mississippi Ground Lease Renewals",
        "category": "Special Purpose Documents",
        "filename": "Form 3759.25 - Mississippi Ground Lease Renewals FNMA Only 07_2021.pdf",
        "dir": "special_purpose",
        "states": ["Mississippi"],
    },
}


def get_additional_docs_for_state(state: str) -> list[dict]:
    """Return additional documents available for a given state."""
    result = []
    for doc_id, doc in ADDITIONAL_DOCUMENTS.items():
        if doc["states"] is None or state in doc["states"]:
            result.append({
                "id": doc_id,
                "name": doc["name"],
                "category": doc["category"],
            })
    return result


# -----------------------------------------------------------------------------
# Sample data generator
# -----------------------------------------------------------------------------
FIRST_NAMES = ["John", "Mary", "Robert", "Jennifer", "Michael", "Linda", "David", "Patricia"]
LAST_NAMES  = ["Smith", "Johnson", "Williams", "Brown", "Jones", "Garcia", "Miller", "Davis"]
LENDERS     = ["Acme Home Lending, Inc.", "Summit Mortgage Corp.", "BlueRidge Financial, LLC",
               "Pioneer Federal Savings Bank", "Evergreen Home Loans"]
TRUSTEES    = ["First American Title Insurance Co.", "Chicago Title Company",
               "Old Republic National Title", "Stewart Title Guaranty Co."]
CITIES_BY_STATE = {
    "California": ["Los Angeles", "San Diego", "Sacramento"],
    "Texas": ["Houston", "Dallas", "Austin"],
    "New York": ["New York", "Buffalo", "Rochester"],
    "Florida": ["Miami", "Orlando", "Tampa"],
}
BORROWER_ADDRESSES = [
    "456 Sunset Blvd, Apt 12, Beverly Hills, CA 90210",
    "789 River Rd, Suite 3B, Brooklyn, NY 11201",
    "321 Lakewood Dr, Miami, FL 33101",
    "1010 Congress Ave, Austin, TX 78701",
]

def make_sample_data(state: str) -> dict:
    """Build one consistent set of example loan data."""
    today = date.today()
    first_payment = (today.replace(day=1) + timedelta(days=45)).replace(day=1)
    maturity = first_payment.replace(year=first_payment.year + 30)

    loan_amount = random.choice([185000, 225000, 320000, 410000, 525000])
    rate = round(random.uniform(5.25, 7.75), 3)
    city = random.choice(CITIES_BY_STATE.get(state, ["Springfield", "Riverside", "Madison"]))
    zip_code = f"{random.randint(10000, 99999)}"
    street = (f"{random.randint(100, 9999)} "
              f"{random.choice(['Maple','Oak','Pine','Cedar','Elm'])} "
              f"{random.choice(['St','Ave','Rd','Ln','Blvd'])}")

    borrower = f"{random.choice(FIRST_NAMES)} {random.choice(LAST_NAMES)}"
    co_borrower = f"{random.choice(FIRST_NAMES)} {random.choice(LAST_NAMES)}"
    borrower_addr = random.choice(BORROWER_ADDRESSES)

    return {
        "borrower_name": borrower,
        "co_borrower_name": co_borrower,
        "borrower_address": borrower_addr,
        "loan_date": today.strftime("%B %d, %Y"),
        "loan_date_month_day": today.strftime("%B %d"),
        "loan_date_year": today.strftime("%Y"),
        "lender_name": random.choice(LENDERS),
        "lender_org_type": "corporation",
        "lender_org_state": "Delaware",
        "lender_address": "100 Finance Plaza, Wilmington, DE 19801",
        "trustee_name": random.choice(TRUSTEES),
        "trustee_address": "200 Trust Center Drive, Los Angeles, CA 90017",
        "loan_amount_number": f"{loan_amount:,}.00",
        "loan_amount_words": _num_to_words(loan_amount),
        "loan_amount_raw": str(loan_amount),
        "interest_rate": f"{rate}",
        "monthly_payment": f"{round(_monthly_pi(loan_amount, rate, 360), 2):,.2f}",
        "first_payment_day": "1st",
        "first_payment_date": first_payment.strftime("%B 1, %Y"),
        "first_payment_month_year": first_payment.strftime("%B, %Y"),
        "maturity_date": maturity.strftime("%B 1, %Y"),
        "maturity_date_month_day": maturity.strftime("%B 1"),
        "maturity_date_year": maturity.strftime("%Y"),
        "late_charge_days": "15",
        "late_charge_percent": "5",
        "property_street": street,
        "property_city": city,
        "property_zip": zip_code,
        "property_county": f"{city} County",
        "property_full_address": f"{street}, {city}, {state} {zip_code}",
        "recording_jurisdiction_type": "County",
        "recording_jurisdiction_name": f"{city} County",
        "state": state,
        "state_abbr": STATE_ABBR.get(state, ""),
        "note_city": city,
        "note_state": state,
        "cancel_lender_name": random.choice(LENDERS),
        "cancel_lender_address": "100 Finance Plaza, Wilmington, DE 19801",
        "cancel_deadline_month_day": (today + timedelta(days=3)).strftime("%B %d"),
        "cancel_deadline_year": (today + timedelta(days=3)).strftime("%Y"),
    }


def _monthly_pi(principal, annual_rate_pct, n_months):
    r = (annual_rate_pct / 100) / 12
    if r == 0:
        return principal / n_months
    return principal * (r * (1 + r) ** n_months) / ((1 + r) ** n_months - 1)


def _num_to_words(n):
    """Very small helper; good enough for sample data."""
    ones = ["", "One", "Two", "Three", "Four", "Five", "Six", "Seven", "Eight", "Nine",
            "Ten", "Eleven", "Twelve", "Thirteen", "Fourteen", "Fifteen", "Sixteen",
            "Seventeen", "Eighteen", "Nineteen"]
    tens = ["", "", "Twenty", "Thirty", "Forty", "Fifty", "Sixty", "Seventy", "Eighty", "Ninety"]
    def below_thousand(x):
        parts = []
        if x >= 100:
            parts.append(ones[x // 100] + " Hundred")
            x %= 100
        if x >= 20:
            t = tens[x // 10]
            if x % 10:
                t += "-" + ones[x % 10]
            parts.append(t)
        elif x > 0:
            parts.append(ones[x])
        return " ".join(parts)
    if n == 0:
        return "Zero"
    result = []
    if n >= 1_000_000:
        result.append(below_thousand(n // 1_000_000) + " Million"); n %= 1_000_000
    if n >= 1_000:
        result.append(below_thousand(n // 1_000) + " Thousand"); n %= 1_000
    if n > 0:
        result.append(below_thousand(n))
    return " ".join(result)


# =============================================================================
# DOCX placeholder filling
# =============================================================================

def _replace_underscores_in_run(run, value: str):
    """Replace underscore blanks in a single run's text with `value`.

    Preserves the run's formatting (bold, italic, font, size, etc.).
    Only replaces contiguous sequences of 3+ underscores.
    """
    run.text = re.sub(r'_{3,}', value, run.text, count=1)


def _fill_paragraph_contextual(paragraph, rules: list[tuple[str, str, str | None]]):
    """Fill underscore blanks in a paragraph using contextual rules.

    Each rule is a tuple: (before_context, value, after_context)
      - before_context: regex pattern for text before the underscore blank
      - value: the replacement text
      - after_context: optional regex pattern for text after the blank (can be None)

    Works at the full paragraph text level first to decide WHICH rule applies,
    then fills at the run level to preserve formatting.
    """
    full_text = paragraph.text

    # Skip paragraphs with no underscore blanks
    if not re.search(r'_{3,}', full_text):
        return

    for before_ctx, value, after_ctx in rules:
        # Build a pattern that matches: before_context + underscores + after_context
        if after_ctx:
            pattern = f'({re.escape(before_ctx)})' + r'\s*(_{3,})' + f'({re.escape(after_ctx)})'
        else:
            pattern = f'({re.escape(before_ctx)})' + r'\s*(_{3,})'

        if re.search(pattern, full_text, re.IGNORECASE):
            # This rule matches. Now replace the underscores in the runs.
            # We need to find the right run(s) that contain this underscore.
            _replace_underscores_in_runs_for_context(paragraph, before_ctx, value)
            # Update full_text for next rule check
            full_text = paragraph.text


def _replace_underscores_in_runs_for_context(paragraph, before_context: str, value: str):
    """Find the right underscore blank near `before_context` and replace it.

    Because runs can split text arbitrarily, we work with the full paragraph
    text to find the position, then map back to runs.
    """
    runs = paragraph.runs
    if not runs:
        return

    # Build cumulative text positions for each run
    positions = []
    offset = 0
    for run in runs:
        positions.append((offset, offset + len(run.text), run))
        offset += len(run.text)

    full_text = paragraph.text

    # Find the underscore position that follows the before_context
    ctx_lower = before_context.lower()
    text_lower = full_text.lower()

    ctx_idx = text_lower.find(ctx_lower)
    if ctx_idx < 0:
        return

    # Find the first underscore blank after the context
    search_start = ctx_idx + len(before_context)
    match = re.search(r'_{3,}', full_text[search_start:])
    if not match:
        return

    blank_start = search_start + match.start()
    blank_end = search_start + match.end()

    # Now replace in the runs that overlap with [blank_start, blank_end)
    replaced = False
    for start_pos, end_pos, run in positions:
        if start_pos >= blank_end:
            break
        if end_pos <= blank_start:
            continue

        # This run overlaps with the blank
        run_text = run.text
        # Find underscores within this run
        local_blank_start = max(0, blank_start - start_pos)
        local_blank_end = min(len(run_text), blank_end - start_pos)

        if not replaced:
            # Replace the underscore portion with the value
            run.text = run_text[:local_blank_start] + value + run_text[local_blank_end:]
            replaced = True
        else:
            # Subsequent runs that were part of the same underscore sequence
            # Remove the underscore portion
            run.text = run_text[:local_blank_start] + run_text[local_blank_end:]


def fill_security_instrument(docx_path: Path, data: dict, state: str) -> Path:
    """Fill placeholders in a Security Instrument DOCX.

    Fills loan/lender/property fields. Leaves borrower/notary fields empty.
    Returns path to the filled DOCX in a temp directory.
    """
    doc = DocxDocument(str(docx_path))

    is_dot = state in DEED_OF_TRUST_STATES

    # ----- Contextual rules for filling -----
    # Each rule: (text_before_blank, replacement_value, text_after_blank_or_None)
    # Rules are applied per-paragraph, so order matters within a paragraph.

    for para in doc.paragraphs:
        text = para.text

        # Skip empty paragraphs
        if not text.strip():
            continue

        # ----- SKIP: Signature lines with (Seal) — leave blank for borrower -----
        if '(Seal)' in text:
            continue

        # ----- SKIP: Witness signature lines -----
        if text.strip() == 'Witnesses:' or text.strip().startswith('Name:____'):
            continue

        # ----- SKIP: Acknowledgment / Notary seal area -----
        if 'Space Below This Line For Acknowledgment' in text or 'Space Above This Line' in text:
            continue

        # ----- (A) Borrower section -----
        if ('\u201cBorrower\u201d is' in text or '"Borrower" is' in text) and 'currently residing' in text:
            rules = []
            # "Borrower" is __________, currently residing at __________
            if '\u201cBorrower\u201d is' in text:
                rules.append(('\u201cBorrower\u201d is ', data.get("borrower_name", ""), None))
            elif '"Borrower" is' in text:
                rules.append(('"Borrower" is ', data.get("borrower_name", ""), None))
            # currently residing at ________
            if 'residing at' in text:
                rules.append(('residing at ', data.get("borrower_address", ""), None))
            _fill_paragraph_contextual(para, rules)
            continue

        # ----- (B) Lender section -----
        if '\u201cLender\u201d is' in text or '"Lender" is' in text:
            rules = []
            # "Lender" is __________.
            if '\u201cLender\u201d is' in text:
                rules.append(('\u201cLender\u201d is ', data["lender_name"], None))
            else:
                rules.append(('"Lender" is ', data["lender_name"], None))
            # Lender is a ______ organized
            if 'is a ' in text and 'organized' in text:
                rules.append(('is a ', data["lender_org_type"], None))
            # under the laws of ______
            if 'under the laws of' in text:
                rules.append(('under the laws of', data["lender_org_state"], None))
                # Some states have "under the laws of______" no space before underscore
                rules.append(('under the laws of ', data["lender_org_state"], None))
            # Lender's address is ______
            if "Lender\u2019s address is" in text or "Lender's address is" in text:
                if "Lender\u2019s address is" in text:
                    rules.append(("Lender\u2019s address is ", data["lender_address"], None))
                else:
                    rules.append(("Lender's address is ", data["lender_address"], None))
            _fill_paragraph_contextual(para, rules)
            continue

        # ----- (C) Trustee section (only for Deed of Trust states) -----
        if is_dot and ('\u201cTrustee\u201d is' in text or '"Trustee" is' in text):
            rules = []
            if '\u201cTrustee\u201d is' in text:
                rules.append(('\u201cTrustee\u201d is ', data["trustee_name"], None))
            else:
                rules.append(('"Trustee" is ', data["trustee_name"], None))
            if "Trustee\u2019s address is" in text or "Trustee's address is" in text:
                if "Trustee\u2019s address is" in text:
                    rules.append(("Trustee\u2019s address is ", data["trustee_address"], None))
                else:
                    rules.append(("Trustee's address is ", data["trustee_address"], None))
            _fill_paragraph_contextual(para, rules)
            continue

        # ----- (D) Note section -----
        if ('\u201cNote\u201d means' in text or '"Note" means' in text):
            rules = []
            # promissory note dated __________, ____
            if 'note dated' in text.lower():
                rules.append(('note dated ', data["loan_date_month_day"], None))
                # year after the comma
                rules.append((data["loan_date_month_day"] + ', ', data["loan_date_year"], None))
            # pay Lender ______ Dollars (U.S. $______)
            if 'pay Lender' in text or 'pay' in text:
                if 'Dollars' in text:
                    rules.append(('pay Lender ', data["loan_amount_words"], None))
                    rules.append(('$', data["loan_amount_number"], None))
            # not later than _______, ______
            if 'not later than' in text:
                rules.append(('not later than', data["maturity_date_month_day"], None))
                # year after last comma
                if re.search(r'not later than.*_{3,}.*,\s*_{3,}', text):
                    rules.append((data["maturity_date_month_day"] + ',', data["maturity_date_year"], None))
            _fill_paragraph_contextual(para, rules)
            continue

        # ----- (F) Security Instrument date -----
        if ('\u201cSecurity Instrument\u201d means' in text or '"Security Instrument" means' in text):
            rules = []
            if 'which is dated' in text:
                rules.append(('which is dated ', data["loan_date_month_day"], None))
                # year
                rules.append((data["loan_date_month_day"] + ',', data["loan_date_year"], None))
            _fill_paragraph_contextual(para, rules)
            continue

        # ----- Transfer of Rights: property location -----
        if 'located in the' in text.lower() and re.search(r'_{3,}', text):
            rules = [
                ('located in the ', data["recording_jurisdiction_type"], None),
            ]
            if ' of ' in text:
                rules.append((' of ', data["recording_jurisdiction_name"], None))
            _fill_paragraph_contextual(para, rules)
            continue

        # ----- Property address -----
        if 'currently has the address of' in text.lower():
            rules = [('address of ', data["property_street"], None)]
            _fill_paragraph_contextual(para, rules)
            continue

        # City, State, Zip line (e.g., " _______, California ______ ")
        state_name = data["state"]
        if state_name in text and re.search(r'_{3,}.*' + re.escape(state_name), text):
            rules = []
            # City blank before the state name
            match = re.search(r'(_{3,})\s*,?\s*' + re.escape(state_name), text)
            if match:
                # Find what's before these underscores
                pre = text[:match.start()].rstrip()
                rules.append((pre[-3:] if pre else '', data["property_city"], None))
            # Zip blank after the state name
            match2 = re.search(re.escape(state_name) + r'\s+(_{3,})', text)
            if match2:
                rules.append((state_name + ' ', data["property_zip"], None))
            if rules:
                _fill_paragraph_contextual(para, rules)
            continue

    # Save the filled document to a temp file
    out_path = CACHE / f"filled_SI_{STATE_ABBR.get(state, state)}.docx"
    doc.save(str(out_path))
    return out_path


def fill_note(note_stem: str, data: dict, state: str) -> Path:
    """Fill placeholders in a Note DOCX.

    Uses the .docx version of the note. Fills loan fields, leaves borrower
    signature fields empty.
    Returns path to the filled DOCX.
    """
    docx_path = NOTE_DIR / (note_stem + ".docx")
    if not docx_path.exists():
        # Try without parens variant
        raise FileNotFoundError(f"Note DOCX not found: {docx_path}")

    doc = DocxDocument(str(docx_path))

    for para in doc.paragraphs:
        text = para.text

        if not text.strip() or '_{3,}' not in text and '___' not in text:
            if '___' not in text:
                continue

        # ----- SKIP: Signature lines with (Seal) — leave blank for borrower -----
        if '(Seal)' in text:
            continue

        # ----- Header line: Date, City, State -----
        # Pattern: "_________, ________  _____________,  ______________"
        # [Note Date]  [City]  [State]
        if '[Note Date]' in text or '[note date]' in text.lower():
            # This is the label line, skip it
            continue

        # The actual date/city/state line has 4 underscore groups
        underscore_groups = list(re.finditer(r'_{3,}', text))
        if len(underscore_groups) >= 3 and para.text.strip().startswith('_'):
            # Header line: Date______, Year____ City______, State______
            rules = []
            # We replace them positionally: date, year, city, state
            _replace_blanks_positionally(para, [
                data["loan_date_month_day"],
                data["loan_date_year"],
                data["note_city"],
                data["note_state"],
            ])
            continue

        # ----- Property Address line -----
        if '[Property Address]' in text:
            continue
        # The line before [Property Address] has the full address
        if text.strip().startswith('___') and len(underscore_groups) == 1 and len(text.strip()) > 50:
            # Full-width underscore line = property address
            _replace_blanks_positionally(para, [data["property_full_address"]])
            continue

        # ----- Section 1: Borrower's Promise to Pay -----
        if 'U.S. $' in text and 'Principal' in text and 'Lender' in text:
            rules = []
            # loan amount: U.S. $ _________
            rules.append(('$ ', data["loan_amount_number"], None))
            # lender name: from ____________ (the "Lender")
            rules.append(('from ', data["lender_name"], None))
            _fill_paragraph_contextual(para, rules)
            continue

        # ----- Section 2: Interest -----
        if 'yearly rate of' in text.lower():
            rules = [('rate of ', data["interest_rate"], None)]
            _fill_paragraph_contextual(para, rules)
            continue

        # ----- Section 3: Payments -----
        if 'Monthly Payment' in text and 'day of each month' in text:
            rules = []
            # the ______ day
            rules.append(('on the ', data["first_payment_day"], None))
            # beginning on __________, _____
            rules.append(('beginning on ', data["first_payment_date"], None))
            # Handle: if on _________, _____ I still owe
            if 'still owe' in text:
                rules.append(('If, on ', data["maturity_date_month_day"], None))
                rules.append((data["maturity_date_month_day"] + ',', data["maturity_date_year"], None))
            _fill_paragraph_contextual(para, rules)
            continue

        # Monthly Payments at (lender address)
        if 'Monthly Payments at' in text:
            rules = [('Payments at ', data["lender_address"], None)]
            _fill_paragraph_contextual(para, rules)
            continue

        # Monthly Payment amount
        if 'Monthly Payment will be in the amount' in text and 'U.S. $' in text:
            rules = [('$', data["monthly_payment"], None)]
            _fill_paragraph_contextual(para, rules)
            continue

        # ----- Section 6: Late Charges -----
        if 'calendar days after' in text:
            rules = [('end of ', data["late_charge_days"], None)]
            _fill_paragraph_contextual(para, rules)
            # charge percent
            if '% of my overdue' in text or '%' in text:
                rules2 = [('will be ', data["late_charge_percent"], None)]
                _fill_paragraph_contextual(para, rules2)
            continue

    # Save
    out_path = CACHE / f"filled_Note_{STATE_ABBR.get(state, state)}.docx"
    doc.save(str(out_path))
    return out_path


def _replace_blanks_positionally(paragraph, values: list[str]):
    """Replace underscore blanks in a paragraph positionally (1st blank = values[0], etc.).

    Handles blanks that span multiple runs by working at the full-text level
    and mapping character positions back to individual runs.
    """
    runs = paragraph.runs
    if not runs:
        return

    # Build cumulative text positions for each run
    positions = []
    offset = 0
    for run in runs:
        positions.append((offset, offset + len(run.text), run))
        offset += len(run.text)

    full_text = paragraph.text

    # Find all contiguous underscore groups (3+ underscores, possibly with
    # non-underscore separators within the same "blank" if they're just
    # formatting splits — but we define a blank as _{3,} in full text)
    blanks = list(re.finditer(r'_{3,}', full_text))

    # Replace each blank with corresponding value
    # We must process from right to left to preserve character positions
    replacements = []
    for i, match in enumerate(blanks):
        if i >= len(values):
            break
        replacements.append((match.start(), match.end(), values[i]))

    # Apply replacements from right to left
    for blank_start, blank_end, value in reversed(replacements):
        # Find and update the runs that overlap with this blank
        replaced = False
        for start_pos, end_pos, run in positions:
            if start_pos >= blank_end:
                continue
            if end_pos <= blank_start:
                continue

            run_text = run.text
            local_blank_start = max(0, blank_start - start_pos)
            local_blank_end = min(len(run_text), blank_end - start_pos)

            if not replaced:
                run.text = run_text[:local_blank_start] + value + run_text[local_blank_end:]
                replaced = True
            else:
                run.text = run_text[:local_blank_start] + run_text[local_blank_end:]


# =============================================================================
# DOCX → PDF conversion
# =============================================================================

def _find_libreoffice() -> str | None:
    """Return the LibreOffice/soffice binary path, or None."""
    for name in ("libreoffice", "soffice"):
        if shutil.which(name):
            return name
    mac_path = "/Applications/LibreOffice.app/Contents/MacOS/soffice"
    if os.path.isfile(mac_path) and os.access(mac_path, os.X_OK):
        return mac_path
    return None


def _docx_to_pdf_python(docx_path: Path, out_pdf: Path) -> Path:
    """Pure-Python DOCX-to-PDF fallback using python-docx + reportlab."""
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = DocxDocument(str(docx_path))
    buf = io.BytesIO()
    styles = getSampleStyleSheet()

    body_style = ParagraphStyle(
        "docx_body", parent=styles["Normal"], fontSize=10,
        leading=13, spaceAfter=4, fontName="Helvetica",
    )
    bold_style = ParagraphStyle(
        "docx_bold", parent=body_style, fontName="Helvetica-Bold",
    )
    title_style = ParagraphStyle(
        "docx_title", parent=styles["Title"], fontSize=14,
        spaceAfter=10, fontName="Helvetica-Bold",
    )
    heading_style = ParagraphStyle(
        "docx_heading", parent=styles["Heading2"], fontSize=12,
        spaceBefore=10, spaceAfter=6, fontName="Helvetica-Bold",
    )
    center_style = ParagraphStyle(
        "docx_center", parent=body_style, alignment=1,
    )

    pdf_doc = SimpleDocTemplate(
        buf, pagesize=LETTER,
        leftMargin=0.75 * inch, rightMargin=0.75 * inch,
        topMargin=0.75 * inch, bottomMargin=0.75 * inch,
    )

    story = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            story.append(Spacer(1, 6))
            continue

        text_esc = (text.replace("&", "&amp;")
                       .replace("<", "&lt;")
                       .replace(">", "&gt;"))

        sname = (para.style.name or "").lower()
        if "title" in sname:
            style = title_style
        elif "heading" in sname:
            style = heading_style
        else:
            try:
                if para.alignment == WD_ALIGN_PARAGRAPH.CENTER:
                    style = center_style
                elif para.runs and all(r.bold for r in para.runs):
                    style = bold_style
                else:
                    style = body_style
            except Exception:
                style = body_style

        parts = []
        for run in para.runs:
            r_text = (run.text.replace("&", "&amp;")
                             .replace("<", "&lt;")
                             .replace(">", "&gt;"))
            if run.bold and run.italic:
                parts.append(f"<b><i>{r_text}</i></b>")
            elif run.bold:
                parts.append(f"<b>{r_text}</b>")
            elif run.italic:
                parts.append(f"<i>{r_text}</i>")
            else:
                parts.append(r_text)

        rich_text = "".join(parts) if parts else text_esc
        try:
            story.append(Paragraph(rich_text, style))
        except Exception:
            story.append(Paragraph(text_esc, body_style))

    if not story:
        story.append(Paragraph("(empty document)", body_style))

    pdf_doc.build(story)
    buf.seek(0)
    out_pdf.write_bytes(buf.read())
    return out_pdf


def add_initials_overlay(pdf_path: Path, out_path: Path) -> Path:
    """Stamp an initials footer on every page of a PDF.

    Adds a thin bar at the bottom of each page with:
      Borrower Initials: ________   Co-Borrower Initials: ________
    Uses reportlab to build a transparent overlay per page, then merges
    it with PyPDF2 so original content is untouched.
    """
    from reportlab.pdfgen import canvas as rl_canvas
    from reportlab.lib.pagesizes import LETTER
    from reportlab.lib import colors as rl_colors

    reader = PdfReader(str(pdf_path))
    writer = PdfWriter()

    for page in reader.pages:
        # Determine page dimensions
        media = page.mediabox
        page_w = float(media.width)
        page_h = float(media.height)

        # Build the overlay in memory
        buf = io.BytesIO()
        c = rl_canvas.Canvas(buf, pagesize=(page_w, page_h))

        # --- footer bar ---
        bar_h = 24          # height of the initials bar in points
        bar_y = 18          # distance from bottom of page
        margin = 36         # left/right margin

        # Light grey background strip
        c.setFillColor(rl_colors.HexColor("#f0f2f7"))
        c.setStrokeColor(rl_colors.HexColor("#c8ccd8"))
        c.setLineWidth(0.5)
        c.rect(margin, bar_y, page_w - 2 * margin, bar_h, fill=1, stroke=1)

        # Text
        c.setFillColor(rl_colors.HexColor("#1a1f36"))
        c.setFont("Helvetica", 7.5)

        label_y = bar_y + 8   # vertical centre of text inside bar

        # Left side — Borrower
        c.drawString(margin + 8, label_y, "Borrower Initials:")
        line_x1 = margin + 90
        c.setLineWidth(0.75)
        c.setStrokeColor(rl_colors.HexColor("#1a1f36"))
        c.line(line_x1, label_y, line_x1 + 72, label_y)

        # Right side — Co-Borrower (centred in right half)
        mid = page_w / 2
        c.drawString(mid + 8, label_y, "Co-Borrower Initials:")
        line_x2 = mid + 110
        c.line(line_x2, label_y, line_x2 + 72, label_y)

        c.save()
        buf.seek(0)

        # Merge overlay onto the page
        overlay_reader = PdfReader(buf)
        overlay_page = overlay_reader.pages[0]
        page.merge_page(overlay_page)
        writer.add_page(page)

    with open(str(out_path), "wb") as f:
        writer.write(f)

    return out_path


def docx_to_pdf(docx_path: Path, cache_name: str | None = None) -> Path:
    """Convert a DOCX to PDF. Uses LibreOffice if available, otherwise
    falls back to a pure-Python conversion."""
    if cache_name:
        out_pdf = CACHE / (cache_name + ".pdf")
    else:
        out_pdf = CACHE / (docx_path.stem + ".pdf")

    # Don't use cache for filled documents - always regenerate
    lo = _find_libreoffice()
    if lo:
        with tempfile.TemporaryDirectory() as td:
            subprocess.run(
                [lo, "--headless", "--convert-to", "pdf",
                 "--outdir", td, str(docx_path)],
                check=True, capture_output=True
            )
            produced = next(Path(td).glob("*.pdf"))
            shutil.move(str(produced), str(out_pdf))
    else:
        print(f"  [info] LibreOffice not found – using Python fallback for {docx_path.name}")
        _docx_to_pdf_python(docx_path, out_pdf)

    return out_pdf


# =============================================================================
# Bundle builder
# =============================================================================

def build_bundle(state: str, data: dict | None, mode: str,
                  additional_doc_ids: list[str] | None = None) -> tuple[bytes, str]:
    """Build the closing package PDF.

    mode: 'empty' = blank templates, 'filled' = fill with data
    additional_doc_ids: list of document IDs from ADDITIONAL_DOCUMENTS to append
    """
    if state not in SECURITY_INSTRUMENTS:
        abort(400, f"Unknown state: {state}")

    si_docx_path = SI_DIR / SECURITY_INSTRUMENTS[state]
    note_stem = STATE_SPECIFIC_NOTES.get(state, MULTISTATE_NOTE)

    merger = PdfWriter()

    if mode == "filled" and data:
        # 1. Fill Security Instrument
        print(f"  [fill] Security Instrument for {state}...")
        filled_si_docx = fill_security_instrument(si_docx_path, data, state)
        si_pdf_raw = docx_to_pdf(filled_si_docx, cache_name=f"filled_SI_{STATE_ABBR.get(state, state)}")
        si_pdf = CACHE / f"filled_SI_{STATE_ABBR.get(state, state)}_initialed.pdf"
        print(f"  [initials] Adding initials footer to Security Instrument...")
        add_initials_overlay(si_pdf_raw, si_pdf)

        # 2. Fill Note
        print(f"  [fill] Note for {state}...")
        try:
            filled_note_docx = fill_note(note_stem, data, state)
            note_pdf = docx_to_pdf(filled_note_docx, cache_name=f"filled_Note_{STATE_ABBR.get(state, state)}")
        except FileNotFoundError:
            # Fall back to the original PDF if no DOCX available
            note_pdf_path = NOTE_DIR / (note_stem + ".pdf")
            if note_pdf_path.exists():
                note_pdf = note_pdf_path
            else:
                abort(500, f"Missing Note: {note_stem}")

        # 3. Notice of Right to Cancel (pre-converted PDF)
        if not NOTICE_PDF.exists():
            abort(500, f"Missing Notice PDF: {NOTICE_PDF}")

        # Merge: SI → (CA Acknowledgement) → Note → Notice
        si_reader = PdfReader(str(si_pdf))
        for page in si_reader.pages:
            merger.add_page(page)

        if state == "California" and CA_ACK_PDF.exists():
            print(f"  [add] California Acknowledgement...")
            ca_ack_reader = PdfReader(str(CA_ACK_PDF))
            for page in ca_ack_reader.pages:
                merger.add_page(page)

        note_reader = PdfReader(str(note_pdf))
        for page in note_reader.pages:
            merger.add_page(page)

        notice_reader = PdfReader(str(NOTICE_PDF))
        for page in notice_reader.pages:
            merger.add_page(page)

        # Append additional documents (riders, special purpose docs)
        if additional_doc_ids:
            for doc_id in additional_doc_ids:
                doc_info = ADDITIONAL_DOCUMENTS.get(doc_id)
                if not doc_info:
                    continue
                doc_dir = RIDER_DIR if doc_info["dir"] == "riders" else SP_DIR
                doc_path = doc_dir / doc_info["filename"]
                if doc_path.exists():
                    print(f"  [add] {doc_info['name']}...")
                    doc_reader = PdfReader(str(doc_path))
                    for page in doc_reader.pages:
                        merger.add_page(page)

    else:
        # Empty: just merge blank templates
        # Convert SI docx to PDF
        blank_si_pdf_raw = CACHE / (si_docx_path.stem + ".pdf")
        if not blank_si_pdf_raw.exists():
            blank_si_pdf_raw = docx_to_pdf(si_docx_path)

        blank_si_pdf = CACHE / (si_docx_path.stem + "_initialed.pdf")
        print(f"  [initials] Adding initials footer to Security Instrument...")
        add_initials_overlay(blank_si_pdf_raw, blank_si_pdf)

        note_pdf_path = NOTE_DIR / (note_stem + ".pdf")
        if not note_pdf_path.exists():
            abort(500, f"Missing Note PDF: {note_pdf_path}")

        si_reader = PdfReader(str(blank_si_pdf))
        for page in si_reader.pages:
            merger.add_page(page)

        if state == "California" and CA_ACK_PDF.exists():
            ca_ack_reader = PdfReader(str(CA_ACK_PDF))
            for page in ca_ack_reader.pages:
                merger.add_page(page)

        note_reader = PdfReader(str(note_pdf_path))
        for page in note_reader.pages:
            merger.add_page(page)

        if NOTICE_PDF.exists():
            notice_reader = PdfReader(str(NOTICE_PDF))
            for page in notice_reader.pages:
                merger.add_page(page)

        # Append additional documents (riders, special purpose docs)
        if additional_doc_ids:
            for doc_id in additional_doc_ids:
                doc_info = ADDITIONAL_DOCUMENTS.get(doc_id)
                if not doc_info:
                    continue
                doc_dir = RIDER_DIR if doc_info["dir"] == "riders" else SP_DIR
                doc_path = doc_dir / doc_info["filename"]
                if doc_path.exists():
                    print(f"  [add] {doc_info['name']}...")
                    doc_reader = PdfReader(str(doc_path))
                    for page in doc_reader.pages:
                        merger.add_page(page)

    out = io.BytesIO()
    merger.write(out)
    out.seek(0)

    tag = "Filled" if mode == "filled" else "Empty"
    filename = f"ClosingPackage_{STATE_ABBR.get(state, state)}_{tag}.pdf"
    return out.read(), filename


# =============================================================================
# Flask app & API routes
# =============================================================================

app = Flask(__name__)

# CORS — allow requests from any origin (Vercel frontend)
CORS(app, resources={r"/api/*": {"origins": "*"}})


@app.route("/")
def home():
    """Root endpoint — health check / welcome."""
    return jsonify({
        "service": "Closing Package Generator API",
        "status": "healthy",
        "endpoints": [
            "/api/health",
            "/api/states",
            "/api/sample-data?state=California",
            "POST /api/generate",
        ]
    })


@app.route("/api/health")
def health():
    """Health check endpoint for Cloud Run."""
    return jsonify({"status": "ok"})


@app.route("/api/states")
def states():
    """Return sorted list of states and DOT states."""
    return jsonify({
        "states": sorted(SECURITY_INSTRUMENTS.keys()),
        "dot_states": sorted(DEED_OF_TRUST_STATES),
    })


@app.route("/api/additional-documents")
def additional_documents():
    """Return available additional documents for a given state."""
    state = request.args.get("state", "")
    if not state or state not in SECURITY_INSTRUMENTS:
        return jsonify({"error": "Valid state is required"}), 400
    docs = get_additional_docs_for_state(state)
    return jsonify({"documents": docs})


@app.route("/api/sample-data")
def sample_data():
    """API endpoint to generate sample data for a given state."""
    state = request.args.get("state", "California")
    if state not in SECURITY_INSTRUMENTS:
        return jsonify({"error": f"Unknown state: {state}"}), 400
    data = make_sample_data(state)
    return jsonify(data)


@app.route("/api/generate", methods=["POST"])
def generate():
    """Generate closing package PDF.

    Accepts either form data or JSON body.
    Returns the PDF as a downloadable file.
    """
    # Accept both JSON and form data
    if request.is_json:
        form_data = request.get_json()
    else:
        form_data = request.form.to_dict()

    state = form_data.get("state", "").strip()
    mode = form_data.get("mode", "empty")

    # Parse additional document selections
    additional_doc_ids = form_data.get("additional_documents", [])
    if isinstance(additional_doc_ids, str):
        # If sent as comma-separated string
        additional_doc_ids = [d.strip() for d in additional_doc_ids.split(",") if d.strip()]

    if not state:
        return jsonify({"error": "State is required"}), 400

    if state not in SECURITY_INSTRUMENTS:
        return jsonify({"error": f"Unknown state: {state}"}), 400

    if mode == "filled":
        # Collect form data
        data = {
            "borrower_name": form_data.get("borrower_name", "").strip(),
            "co_borrower_name": form_data.get("co_borrower_name", "").strip(),
            "borrower_address": form_data.get("borrower_address", "").strip(),
            "loan_date": form_data.get("loan_date", "").strip(),
            "lender_name": form_data.get("lender_name", "").strip(),
            "lender_org_type": form_data.get("lender_org_type", "").strip(),
            "lender_org_state": form_data.get("lender_org_state", "").strip(),
            "lender_address": form_data.get("lender_address", "").strip(),
            "trustee_name": form_data.get("trustee_name", "").strip(),
            "trustee_address": form_data.get("trustee_address", "").strip(),
            "loan_amount_number": form_data.get("loan_amount_number", "").strip(),
            "loan_amount_words": form_data.get("loan_amount_words", "").strip(),
            "interest_rate": form_data.get("interest_rate", "").strip(),
            "monthly_payment": form_data.get("monthly_payment", "").strip(),
            "first_payment_day": form_data.get("first_payment_day", "1st").strip(),
            "first_payment_date": form_data.get("first_payment_date", "").strip(),
            "maturity_date": form_data.get("maturity_date", "").strip(),
            "late_charge_days": form_data.get("late_charge_days", "15").strip(),
            "late_charge_percent": form_data.get("late_charge_percent", "5").strip(),
            "property_street": form_data.get("property_street", "").strip(),
            "property_city": form_data.get("property_city", "").strip(),
            "property_zip": form_data.get("property_zip", "").strip(),
            "property_county": form_data.get("property_county", "").strip(),
            "recording_jurisdiction_type": form_data.get("recording_jurisdiction_type", "County").strip(),
            "recording_jurisdiction_name": form_data.get("recording_jurisdiction_name", "").strip(),
            "state": state,
            "state_abbr": STATE_ABBR.get(state, ""),
            "cancel_lender_name": form_data.get("lender_name", "").strip(),
            "cancel_lender_address": form_data.get("lender_address", "").strip(),
            "cancel_deadline_month_day": form_data.get("cancel_deadline_month_day", "").strip(),
            "cancel_deadline_year": form_data.get("cancel_deadline_year", "").strip(),
        }

        # Parse dates into components
        loan_date_str = data["loan_date"]
        if loan_date_str:
            try:
                from datetime import datetime
                ld = datetime.strptime(loan_date_str, "%Y-%m-%d")
                data["loan_date_month_day"] = ld.strftime("%B %d")
                data["loan_date_year"] = ld.strftime("%Y")
                data["loan_date"] = ld.strftime("%B %d, %Y")
            except ValueError:
                data["loan_date_month_day"] = loan_date_str
                data["loan_date_year"] = ""
        else:
            data["loan_date_month_day"] = ""
            data["loan_date_year"] = ""

        maturity_str = data["maturity_date"]
        if maturity_str:
            try:
                from datetime import datetime
                md = datetime.strptime(maturity_str, "%Y-%m-%d")
                data["maturity_date_month_day"] = md.strftime("%B %d")
                data["maturity_date_year"] = md.strftime("%Y")
                data["maturity_date"] = md.strftime("%B %d, %Y")
            except ValueError:
                data["maturity_date_month_day"] = maturity_str
                data["maturity_date_year"] = ""
        else:
            data["maturity_date_month_day"] = ""
            data["maturity_date_year"] = ""

        first_payment_str = data["first_payment_date"]
        if first_payment_str:
            try:
                from datetime import datetime
                fp = datetime.strptime(first_payment_str, "%Y-%m-%d")
                data["first_payment_date"] = fp.strftime("%B %d, %Y")
                data["first_payment_month_year"] = fp.strftime("%B, %Y")
            except ValueError:
                data["first_payment_month_year"] = first_payment_str

        cancel_deadline_str = form_data.get("cancel_deadline", "").strip()
        if cancel_deadline_str:
            try:
                from datetime import datetime
                cd = datetime.strptime(cancel_deadline_str, "%Y-%m-%d")
                data["cancel_deadline_month_day"] = cd.strftime("%B %d")
                data["cancel_deadline_year"] = cd.strftime("%Y")
            except ValueError:
                pass

        # Build property full address
        data["property_full_address"] = (
            f"{data['property_street']}, {data['property_city']}, "
            f"{state} {data['property_zip']}"
        ).strip(", ")

        # Note city/state
        data["note_city"] = data["property_city"]
        data["note_state"] = state

        # Recording jurisdiction defaults
        if not data["recording_jurisdiction_name"]:
            data["recording_jurisdiction_name"] = data["property_county"]

        pdf_bytes, filename = build_bundle(state, data, "filled", additional_doc_ids)
    else:
        pdf_bytes, filename = build_bundle(state, None, "empty", additional_doc_ids)

    return send_file(io.BytesIO(pdf_bytes),
                     mimetype="application/pdf",
                     as_attachment=True,
                     download_name=filename)


if __name__ == "__main__":
    port = int(os.environ.get("PORT", 8080))
    app.run(debug=True, host="0.0.0.0", port=port)
